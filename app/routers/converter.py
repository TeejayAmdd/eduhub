import io
import zipfile

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

import app.models as models
from app.auth import get_current_user

router = APIRouter(prefix="/api/converter", tags=["converter"])

MAX_SIZE = 20 * 1024 * 1024  # 20 MB

SUPPORTED = {
    "pdf_to_docx": {"accepts": [".pdf"], "output_ext": "docx"},
    "pdf_to_txt":  {"accepts": [".pdf"], "output_ext": "txt"},
    "pdf_to_png":  {"accepts": [".pdf"], "output_ext": "png"},
    "image_to_pdf": {"accepts": [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"], "output_ext": "pdf"},
    "docx_to_txt": {"accepts": [".docx"], "output_ext": "txt"},
    "png_to_jpg":  {"accepts": [".png"], "output_ext": "jpg"},
    "jpg_to_png":  {"accepts": [".jpg", ".jpeg"], "output_ext": "png"},
    "image_to_webp": {"accepts": [".jpg", ".jpeg", ".png", ".bmp", ".tiff"], "output_ext": "webp"},
}


@router.post("/convert")
async def convert_document(
    file: UploadFile = File(...),
    conversion_type: str = Form(...),
    current_user: models.User = Depends(get_current_user),
):
    if conversion_type not in SUPPORTED:
        raise HTTPException(400, f"Unsupported conversion: {conversion_type}")

    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(413, "File exceeds 20 MB limit.")

    filename = file.filename or "file"
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""
    if ext not in SUPPORTED[conversion_type]["accepts"]:
        raise HTTPException(400, f"File type {ext} not accepted for this conversion.")

    stem = filename.rsplit(".", 1)[0]

    try:
        if conversion_type == "pdf_to_docx":
            return _pdf_to_docx(content, stem)
        if conversion_type == "pdf_to_txt":
            return _pdf_to_txt(content, stem)
        if conversion_type == "pdf_to_png":
            return _pdf_to_png(content, stem)
        if conversion_type == "image_to_pdf":
            return _image_to_pdf(content, stem)
        if conversion_type == "docx_to_txt":
            return _docx_to_txt(content, stem)
        target = conversion_type.split("_to_")[1]
        return _convert_image_format(content, stem, target)

    except HTTPException:
        raise
    except ImportError as exc:
        raise HTTPException(503, f"Required library not installed: {exc}")
    except Exception as exc:
        raise HTTPException(500, f"Conversion failed: {exc}")


# ── helpers ──────────────────────────────────────────────────────────────────

def _pdf_to_docx(content: bytes, stem: str) -> StreamingResponse:
    import pymupdf
    from docx import Document
    from docx.shared import Pt

    pdf = pymupdf.open(stream=content, filetype="pdf")
    word = Document()

    for page_num, page in enumerate(pdf):
        if page_num > 0:
            word.add_page_break()

        blocks = page.get_text("blocks")  # [(x0,y0,x1,y1,text,block_no,type), ...]
        for block in sorted(blocks, key=lambda b: (b[1], b[0])):
            text = block[4].strip()
            block_type = block[6]
            if block_type == 0 and text:  # 0 = text block
                word.add_paragraph(text)

    pdf.close()

    buf = io.BytesIO()
    word.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{stem}.docx"'},
    )


def _pdf_to_txt(content: bytes, stem: str) -> StreamingResponse:
    import pymupdf

    doc = pymupdf.open(stream=content, filetype="pdf")
    text = "\n\n".join(page.get_text() for page in doc)
    doc.close()

    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{stem}.txt"'},
    )


def _pdf_to_png(content: bytes, stem: str) -> StreamingResponse:
    import pymupdf

    doc = pymupdf.open(stream=content, filetype="pdf")
    mat = pymupdf.Matrix(2, 2)  # 2× zoom for readable resolution

    if len(doc) == 1:
        pix = doc[0].get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        doc.close()
        return StreamingResponse(
            io.BytesIO(img_bytes),
            media_type="image/png",
            headers={"Content-Disposition": f'attachment; filename="{stem}.png"'},
        )

    # Multi-page PDF → ZIP of PNGs
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            zf.writestr(f"{stem}_page_{i + 1:03d}.png", pix.tobytes("png"))
    doc.close()
    zip_buf.seek(0)

    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{stem}_pages.zip"'},
    )


def _image_to_pdf(content: bytes, stem: str) -> StreamingResponse:
    from PIL import Image

    img = Image.open(io.BytesIO(content))
    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")

    buf = io.BytesIO()
    img.save(buf, format="PDF", resolution=100.0)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{stem}.pdf"'},
    )


def _docx_to_txt(content: bytes, stem: str) -> StreamingResponse:
    from docx import Document

    doc = Document(io.BytesIO(content))
    text = "\n".join(para.text for para in doc.paragraphs)

    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{stem}.txt"'},
    )


def _convert_image_format(content: bytes, stem: str, target: str) -> StreamingResponse:
    from PIL import Image

    fmt_map = {
        "jpg":  ("JPEG", "image/jpeg",  "jpg"),
        "png":  ("PNG",  "image/png",   "png"),
        "webp": ("WEBP", "image/webp",  "webp"),
    }
    pil_fmt, mime, ext = fmt_map[target]

    img = Image.open(io.BytesIO(content))
    if pil_fmt == "JPEG" and img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")

    buf = io.BytesIO()
    kwargs = {"quality": 95} if pil_fmt == "JPEG" else {}
    img.save(buf, format=pil_fmt, **kwargs)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{stem}.{ext}"'},
    )
